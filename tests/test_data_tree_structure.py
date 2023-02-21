"""
Test data structure
"""
import os

import filecmp
import difflib
import docx
import pandas as pd
import pytest
from mapping.data_manipulation import init_index, get_data, init_csv, init_txt, init_docx, keep_fr_local
from mapping.utils.node_proposal import NodeProposal
from mapping.utils.node_comment import NodeComment

TEST_PATH = os.path.split(os.path.realpath(__file__))[0]

df_prop_config1, df_coms_config1 = get_data(local_json_file_path=os.path.join(TEST_PATH,
                                                                              "../test_data/test_config_1.json"))
df_prop_config2, df_coms_config2 = get_data(local_json_file_path=os.path.join(TEST_PATH,
                                                                              "../test_data/test_config_2.json"))
df_prop_config3, df_coms_config3 = get_data(local_json_file_path=os.path.join(TEST_PATH,
                                                                              "../test_data/test_config_3.json"))


def nodes_equal(node1, node2):
    """
    This function is used to check if two node objects are equals.
    It is a recursive function that will although go through the children of the
    objects.
    :param node1: first node object
    :param node2: second node object
    :return: return a boolean for each configuration
    """
    is_equal = node1.body == node2.body
    for child_node_1, child_node_2 in zip(node1.children, node2.children):
        is_equal = is_equal and nodes_equal(child_node_1, child_node_2)
    return is_equal


def docx_comparison(path_file1, path_file2):
    """
    Function that will compare to docx file and store the differential in a list.
    :param path_file1:  path to the file created by the program
    :param path_file2: path to the file storing the expected format
    """
    file1 = docx.Document(path_file1)
    file2 = docx.Document(path_file2)
    para1 = ''
    para2 = ''
    for para in file1.paragraphs:
        para1 = para1 + para.text + '\n'
    for para in file2.paragraphs:
        para2 = para2 + para.text + '\n'
    differential = difflib.context_diff(para1.splitlines(), para2.splitlines(), fromfile="test.txt", tofile="test2.txt")
    return list(differential)


CONFIG_1 = [NodeProposal(title="P0 titre",
                         body="P0 body",
                         category="cat",
                         supports=12,
                         nb_comments=4,
                         children=[NodeComment(body="C1",
                                               children=[NodeComment(body="C0",
                                                                     children=[NodeComment(body="R1",
                                                                                           children=[])]),
                                                         NodeComment(body="R2", children=[])])])]

CONFIG_2 = [NodeProposal(title="P0 titre",
                         body="P0",
                         category="cat",
                         supports=12,
                         nb_comments=4,
                         children=[NodeComment(body='C0', children=[]),
                                   NodeComment(body='C1', children=[]),
                                   NodeComment(body='C2', children=[]),
                                   NodeComment(body='C3', children=[])]),
            NodeProposal(title="P1 titre",
                         body="P1",
                         category="cat",
                         supports=5,
                         nb_comments=1,
                         children=[NodeComment(body='C4', children=[])])]

CONFIG_3 = [NodeProposal(title="P0 titre",
                         body="P0",
                         category="cat",
                         supports=12,
                         nb_comments=6,
                         children=[NodeComment(body='C1',
                                               children=[NodeComment(body="R1",
                                                                     children=[NodeComment(body="R2",
                                                                                           children=[NodeComment(body="R3", children=[])])])]),
                                   NodeComment(body='C2', children=[NodeComment(body="R4", children=[])])]),
            NodeProposal(title="P1 titre",
                         body="P1",
                         category="cat",
                         supports=5,
                         nb_comments=0,
                         children=[])]

TEST_CASES_STRUCTURE = [(df_prop_config1, df_coms_config1, CONFIG_1),
                        (df_prop_config2, df_coms_config2, CONFIG_2),
                        (df_prop_config3, df_coms_config3, CONFIG_3)]

TEST_TXT_OUTPUT = [(df_prop_config1, df_coms_config1, os.path.join(TEST_PATH+'/..',
                                                                   "test_data/mapping_result_config1.txt")),
                   (df_prop_config2, df_coms_config2, os.path.join(TEST_PATH+'/..',
                                                                   "test_data/mapping_result_config2.txt")),
                   (df_prop_config3, df_coms_config3, os.path.join(TEST_PATH+'/..',
                                                                   "test_data/mapping_result_config3.txt"))]

TEST_DOCX_OUTPUT = [(df_prop_config1, df_coms_config1, os.path.join(TEST_PATH+"/..",
                                                                    "test_data/mapping_result_config1.docx")),
                    (df_prop_config2, df_coms_config2, os.path.join(TEST_PATH+"/..",
                                                                    "test_data/mapping_result_config2.docx")),
                    (df_prop_config3, df_coms_config3, os.path.join(TEST_PATH+"/..",
                                                                    "test_data/mapping_result_config3.docx"))]


@pytest.mark.parametrize("proposals_dataframe, comments_dataframe, output ", TEST_CASES_STRUCTURE)
def test_tree_structure(proposals_dataframe, comments_dataframe, output):
    """
    This function will test the tree data structure created with init_index
    :param proposals_dataframe: initial dataframe storing the proposals
    :param comments_dataframe: initial dataframe storing the comments
    :param output: handmade configuration used to validate the behavior of the function
    """
    nodes = init_index(proposals_dataframe, comments_dataframe)
    for node_test, node_validation in zip(nodes, output):
        assert nodes_equal(node_test, node_validation)


@pytest.mark.parametrize("proposals_dataframe, comments_dataframe, output ", TEST_CASES_STRUCTURE)
def test_csv_integrity(proposals_dataframe, comments_dataframe, output):
    pytest.skip("TODO: fix this test")
    """
    This function will check the file created with the function init_csv contains the
    adequate number of proposal objects
    :param proposals_dataframe: initial dataframe storing the proposals
    :param comments_dataframe: initial dataframe storing the comments
    :param output: handmade configuration used to validate the behavior of the function
    """
    hash_prop = init_index(proposals_dataframe, comments_dataframe)
    for elem in output:
        if len(elem.children) == 0:
            output.remove(elem)
    df_test = init_csv(hash_prop)
    assert len(df_test)-1 == len(output)


@pytest.mark.parametrize("proposals_dataframe, comments_dataframe, output", TEST_TXT_OUTPUT)
def test_txt_integrity(proposals_dataframe, comments_dataframe, output):
    pytest.skip("TODO: fix this test")
    """
    This function will check that the files created by the function init_txt and verify
    that it is exactly identical to the handmade one
    :param proposals_dataframe: initial dataframe storing the proposals
    :param comments_dataframe: initial dataframe storing the comments
    :param output: handmade configuration used to validate the behavior of the function
    """
    dir_path = os.path.dirname(os.path.realpath(__file__))
    hash_prop = init_index(proposals_dataframe, comments_dataframe)
    init_txt(hash_prop)
    assert filecmp.cmp(os.path.join(dir_path, "../dist/mapping_proposals_comments.txt"), output)


@pytest.mark.parametrize("proposals_dataframe, comments_dataframe, output", TEST_DOCX_OUTPUT)
def test_docx_integrity(proposals_dataframe, comments_dataframe, output):
    pytest.skip("TODO: fix this test")
    """
    This function is used to test that all docx created by the program remains identical
    to the expected configuration
    :param proposals_dataframe: initial dataframe storing the proposals
    :param comments_dataframe: initial dataframe storing the comments
    :param output: handmade configuration used to validate the behavior of the function
    """
    dir_path = os.path.dirname(os.path.realpath(__file__))
    hash_prop = init_index(proposals_dataframe, comments_dataframe)
    init_docx(hash_prop)
    diff = docx_comparison(os.path.join(dir_path, "../dist/mapping_proposals_comments.docx"), output)
    assert len(diff) == 0


def test_keep_fr():
    """
    Test the function responsible of renaming the columns in the dataframe
    used to keep only french locale.
    """
    data = {'body/fr': ["lorem ipsum", "lorem ipsum"],
            "body/en": ["lorem", "lorem"],
            "category/name/fr": ["cat", "cat"],
            "title/fr": ["lorem ipsum", "lorem ipsum"]}
    dataframe = pd.DataFrame(data)
    df_updated = keep_fr_local(dataframe)
    assert set(list(df_updated.columns)) == {'body', 'body/en', "category", 'title'}
