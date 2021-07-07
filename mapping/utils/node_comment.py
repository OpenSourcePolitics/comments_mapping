"""
This file defines the child class NodeProposal which inherits from the class Node
"""
import re
from docx.shared import Inches
from mapping.utils.node import Node


class NodeComment(Node):
    """
    Child class that inherits from the Node class. It is used to manage the information
    specific to comments.
    """

    def __init__(self, body, children):
        Node.__init__(self, body, children)

    def write_txt(self, indent_num, file):
        """
        Recursive function that will iterate on a proposal and all of its children in order
        to write down all the information in a structured way on a txt file
        :param indent_num: initial indentation value
        :type indent_num: int
        :param file: opened file cf- init_txt()
        """
        indent = 2 * indent_num * " "
        file.write(indent + "Commentaire : " + re.sub(r'\n', '\n{}'.format(indent), self.body) + '\n')
        for child in self.children:
            child.write_txt(indent_num + 1, file)

    def get_attributes_as_list(self, node_list):
        """
        Function that will iterate from a proposal through all of its children to
        retrieve the information stored in the object and append it successively to a list
        :param node_list: list that will be progressively filled
        :type node_list: list
        :return: [proposal title, proposal body, comment 1, ..., comment n]
        :rtype: list
        """
        node_list.append(self.body)
        for child in self.children:
            child.get_attributes_as_list(node_list)
        return node_list

    def write_docx(self, indent_num, document):
        """
        This function will append all information stored in the comment objects in a .docx file.
        It will then iterate on the children to retrieve the comments
        :param indent_num: indentation system for the structure
        :type indent_num: float
        :param document: document object
        """
        indent_num = 2*indent_num
        paragraph = document.add_paragraph()
        paragraph.add_run("Commentaire : ").bold = True
        paragraph.add_run(str(self.body) + '\n')
        paragraph.paragraph_format.left_indent = Inches(indent_num)
        for child in self.children:
            child.write_docx(indent_num + 0.05, document)
