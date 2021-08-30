"""
This file defines the child class NodeProposal which inherits from the class Node
"""
from docx.shared import Pt
from mapping.utils.node import Node


class NodeProposal(Node):
    """
    Child class that inherits from the Node class. It is used to manage the information
    specific to proposals.
    """

    # pylint: disable=too-many-arguments
    def __init__(self, title, body, children, supports, nb_comments, category):
        Node.__init__(self, body, children)
        self.title = title
        self.supports = supports
        self.nb_comments = nb_comments
        self.category = category
    # pylint: enable=too-many-arguments

    def write_txt(self, indent_num, file):
        """
        Recursive function that will iterate on a proposal and all of its children in order
        to write down all the information in a structured way on a txt file
        :param indent_num: initial indentation value
        :type indent_num: int
        :param file: opened file cf- init_txt()
        """
        indent = 2 * indent_num * " "
        file.write(indent + 'Title : ' + str(self.title) + '\n')
        file.write(indent + 'Body : ' + str(self.body) + '\n')
        file.write(indent + 'Category : ' + str(self.category) + '\n')
        file.write(indent + 'Supports : ' + str(self.supports) + '\n')
        file.write(indent + 'Amount of comments : ' + str(self.nb_comments) + '\n\n')
        for child in self.children:
            child.write_txt(indent_num + 1, file)

    def get_attributes_as_list(self, node_list):
        """
        Function that will iterate from a proposal through all of its children to
        retrieve the information stored in the object and append it successively to a list
        :param node_list: list that will be progressively filled
        :type node_list: list
        :return: [proposal title, proposal body, comment 1, ..., comment n]
        :rtype: list
        """
        node_list.append(self.title)
        node_list.append(self.body)
        node_list.append(self.category)
        node_list.append(self.supports)
        node_list.append(self.nb_comments)
        for child in self.children:
            child.get_attributes_as_list(node_list)
        return node_list

    def write_docx(self, indent_num, document):
        """
        This function will append all information stored in the proposal objects in a .docx file.
        It will then iterate on the children to retrieve the comments
        :param indent_num: indentation system for the structure
        :type indent_num: float
        :param document: document object
        """
        style = document.styles['Normal']
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)

        title = document.add_heading(self.title, 2)
        title.alignment = 1

        paragraph = document.add_paragraph(self.body + '\n')
        paragraph.style = document.styles['Normal']

        paragraph.add_run("Category : ").bold = True
        paragraph.add_run(str(self.category) + '\n')
        paragraph.add_run("Supports : ").bold = True
        paragraph.add_run(str(self.supports) + '\n')
        paragraph.add_run("Amount of comments : ").bold = True
        paragraph.add_run(str(self.nb_comments) + '\n')

        for child in self.children:
            child.write_docx(indent_num + 0.05, document)
